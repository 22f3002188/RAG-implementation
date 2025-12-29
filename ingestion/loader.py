import os
from langchain_core.documents import Document
from docx import Document as DocxDocument
from PIL import Image
import pytesseract

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


def load_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return Document(page_content=f.read(), metadata={"source": path})


def load_docx(path):
    doc = DocxDocument(path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    if text.strip():
        return Document(page_content=text, metadata={"source": path})
    return None


def load_image(path):
    text = pytesseract.image_to_string(Image.open(path))
    if text.strip():
        return Document(page_content=text, metadata={"source": path})
    return None


def load_files(directory):
    documents = []

    for file in os.listdir(directory):
        path = os.path.join(directory, file)

        if file.endswith(".txt"):
            documents.append(load_txt(path))

        elif file.endswith(".docx"):
            doc = load_docx(path)
            if doc:
                documents.append(doc)

        elif file.lower().endswith((".png", ".jpg", ".jpeg")):
            doc = load_image(path)
            if doc:
                documents.append(doc)

    # remove None entries
    return [d for d in documents if d]
